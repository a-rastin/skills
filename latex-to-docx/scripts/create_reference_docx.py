#!/usr/bin/env python3
"""Create a Persian-aware reference DOCX for Pandoc."""
from __future__ import annotations

import argparse
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_style_fonts(style, persian_font: str, latin_font: str, size_pt: float | None = None, bold: bool | None = None):
    style.font.name = latin_font
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    rPr = style.element.get_or_add_rPr()
    rFonts = ensure_child(rPr, "w:rFonts")
    for attr, value in {
        "w:ascii": latin_font,
        "w:hAnsi": latin_font,
        "w:eastAsia": persian_font,
        "w:cs": persian_font,
    }.items():
        rFonts.set(qn(attr), value)
    lang = ensure_child(rPr, "w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "en-US")
    lang.set(qn("w:bidi"), "fa-IR")


def set_style_rtl(style, alignment: WD_ALIGN_PARAGRAPH | None = None):
    pPr = style.element.get_or_add_pPr()
    bidi = ensure_child(pPr, "w:bidi")
    bidi.set(qn("w:val"), "1")
    if alignment is not None:
        jc = ensure_child(pPr, "w:jc")
        jc.set(qn("w:val"), {
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "both",
            WD_ALIGN_PARAGRAPH.LEFT: "left",
        }.get(alignment, "right"))


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    for style in doc.styles:
        if style.name == name or style.style_id == name:
            return style
    return doc.styles.add_style(name, style_type)


def build_reference(output: Path, persian_font: str, latin_font: str, mono_font: str, paper: str, margins_cm: float):
    pandoc = shutil.which("pandoc")
    if pandoc:
        with tempfile.NamedTemporaryFile(suffix=".docx") as tmp:
            proc = subprocess.run([pandoc, "--print-default-data-file", "reference.docx"], stdout=tmp, stderr=subprocess.PIPE, check=False)
            tmp.flush()
            if proc.returncode == 0 and Path(tmp.name).stat().st_size > 0:
                doc = Document(tmp.name)
            else:
                doc = Document()
    else:
        doc = Document()
    section = doc.sections[0]
    if paper.lower() == "a4":
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    section.top_margin = Cm(margins_cm)
    section.bottom_margin = Cm(margins_cm)
    section.left_margin = Cm(margins_cm)
    section.right_margin = Cm(margins_cm)

    body_styles = ["Normal", "Body Text", "First Paragraph", "Bibliography", "Footnote Text"]
    for name in body_styles:
        style = ensure_style(doc, name)
        set_style_fonts(style, persian_font, latin_font, 12)
        set_style_rtl(style, WD_ALIGN_PARAGRAPH.JUSTIFY if name in {"Normal", "Body Text", "First Paragraph"} else WD_ALIGN_PARAGRAPH.RIGHT)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.15

    heading_sizes = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 12}
    for level, size in heading_sizes.items():
        style = ensure_style(doc, f"Heading {level}")
        set_style_fonts(style, persian_font, latin_font, size, bold=True)
        set_style_rtl(style, WD_ALIGN_PARAGRAPH.RIGHT)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    for name, size, alignment in [
        ("Title", 22, WD_ALIGN_PARAGRAPH.CENTER),
        ("Subtitle", 14, WD_ALIGN_PARAGRAPH.CENTER),
        ("Author", 12, WD_ALIGN_PARAGRAPH.CENTER),
        ("Date", 11, WD_ALIGN_PARAGRAPH.CENTER),
        ("Caption", 10, WD_ALIGN_PARAGRAPH.CENTER),
    ]:
        style = ensure_style(doc, name)
        set_style_fonts(style, persian_font, latin_font, size, bold=name == "Title")
        set_style_rtl(style, alignment)
        style.paragraph_format.keep_with_next = name == "Caption"

    for name in ("Figure Caption", "Table Caption"):
        style = ensure_style(doc, name)
        set_style_fonts(style, persian_font, latin_font, 10)
        set_style_rtl(style, WD_ALIGN_PARAGRAPH.CENTER)
        style.paragraph_format.keep_with_next = True

    for name in ("Source Code", "Code Block", "Verbatim Char"):
        style_type = WD_STYLE_TYPE.CHARACTER if name == "Verbatim Char" else WD_STYLE_TYPE.PARAGRAPH
        style = ensure_style(doc, name, style_type)
        set_style_fonts(style, mono_font, mono_font, 9)
        if style_type == WD_STYLE_TYPE.PARAGRAPH:
            set_style_rtl(style, WD_ALIGN_PARAGRAPH.LEFT)

    compact = ensure_style(doc, "Compact")
    set_style_fonts(compact, persian_font, latin_font, 10)
    compact.paragraph_format.space_after = Pt(0)
    compact.paragraph_format.line_spacing = 1.0
    table_style = ensure_style(doc, "Table", WD_STYLE_TYPE.TABLE)
    set_style_fonts(table_style, persian_font, latin_font, 10)

    settings = doc.settings.element
    theme = settings.find(qn("w:themeFontLang"))
    if theme is None:
        theme = OxmlElement("w:themeFontLang")
        settings.append(theme)
    theme.set(qn("w:val"), "en-US")
    theme.set(qn("w:eastAsia"), "en-US")
    theme.set(qn("w:bidi"), "fa-IR")
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    # Pandoc ignores body content in a reference DOCX. Keep one empty paragraph.
    if not doc.paragraphs:
        doc.add_paragraph("")
    else:
        doc.paragraphs[0].text = ""
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--persian-font", default="Noto Naskh Arabic")
    parser.add_argument("--latin-font", default="Liberation Serif")
    parser.add_argument("--mono-font", default="DejaVu Sans Mono")
    parser.add_argument("--paper", default="A4")
    parser.add_argument("--margins-cm", type=float, default=2.5)
    args = parser.parse_args()
    build_reference(args.output, args.persian_font, args.latin_font, args.mono_font, args.paper, args.margins_cm)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
