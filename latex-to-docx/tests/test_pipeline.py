from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from inspect_latex import inventory
from postprocess_docx import process_docx
from validate_docx import validate


class PipelineTests(unittest.TestCase):
    def test_inventory_and_include_graph(self):
        inv = inventory(ROOT / "tests" / "fixtures" / "main.tex")
        self.assertEqual(inv["counts"]["source_files"], 2)
        self.assertGreater(inv["persian_signals"]["arabic_character_count"], 0)
        self.assertEqual(inv["counts"]["tables"], 1)
        self.assertGreaterEqual(inv["counts"]["math_display_estimate"], 1)
        self.assertFalse(inv["missing_dependencies"])

    def test_rtl_postprocess_and_validation(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            original = td / "original.docx"
            patched = td / "patched.docx"
            doc = Document()
            doc.add_paragraph("متن فارسی با نیم‌فاصله و English")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "مقدار"
            table.cell(0, 1).text = "Value"
            doc.save(original)
            process_docx(original, patched)
            report = validate(patched)
            self.assertTrue(report["valid"], json.dumps(report, ensure_ascii=False, indent=2))
            with zipfile.ZipFile(patched) as zf:
                xml = zf.read("word/document.xml").decode("utf-8")
                self.assertIn("w:bidi", xml)
                self.assertIn("w:rtl", xml)
                self.assertIn("w:bidiVisual", xml)

    def test_reference_doc_creation_cli(self):
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "reference.docx"
            proc = subprocess.run([sys.executable, str(SCRIPTS / "create_reference_docx.py"), "--output", str(out)], capture_output=True, text=True)
            self.assertEqual(proc.returncode, 0, proc.stderr)
            self.assertTrue(out.exists())
            with zipfile.ZipFile(out) as zf:
                self.assertIn("word/styles.xml", zf.namelist())
                styles = zf.read("word/styles.xml").decode("utf-8")
                self.assertIn('w:styleId="Table"', styles)
                self.assertIn('w:styleId="Compact"', styles)


if __name__ == "__main__":
    unittest.main()
